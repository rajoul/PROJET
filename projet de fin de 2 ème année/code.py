from docx import Document
from nltk import wordpunct_tokenize
from tkinter import *
from tkinter import filedialog
import pandas as pd
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
import io,re,time
import shutil
import os,subprocess
score_total=0

fenetre=Tk()
fenetre.geometry('600x350')
fenetre.title("PFA")

Texte0=StringVar()
Texte1=StringVar()
Texte2=StringVar()
Texte3=StringVar()
Texte4=StringVar()
Texte5=StringVar()
Texte6=StringVar()
Texte7=StringVar()
Texte8=StringVar()
Texte9=StringVar()
tab=[Texte1,Texte3,Texte5,Texte7]
tab1=[Texte0,Texte2,Texte4,Texte6]

#https://realpython.com/python-keras-text-classification/
principale=["CERTIFICATION AUTHORITIES",'REGISTRATION AUTHORITIES ','VALIDATION AUTHORITIES ',
          'CERTIFICATE ACCEPTANCE']
other=['KEY SIZE','AUTHENTICATION OF ORGANIZATION IDENTITY']
word={}
scoree=0
#f=subprocess.call('lowriter --invisible --convert-to doc "{}"'
                            #.format('file.pdf'), shell=True)
def model_train(word):
    score_total=0
    valeur=0
    for key in word.keys():
        list = []
        fichier = key.split(' ')
        file = fichier[0] + '_' + fichier[1] + '.txt'
        dat = pd.read_csv(file, names=['sentence'], sep='_')
        list.append(dat)
        df = pd.concat(list)
        vectorizer = CountVectorizer(min_df=0, lowercase=False)
        vectorizer.fit([word[key]])
        y = [i for i in range(len(df['sentence'].values))]

        d = vectorizer.transform([word[key]]).toarray()
        sentences_train, sentences_test, y_train, y_test = train_test_split(df['sentence'].values, y, test_size=0.75,random_state=1000)

        vectorizer.fit(sentences_train)
        X_train = vectorizer.transform(sentences_train)
        X_test = vectorizer.transform(sentences_test)

        classifier = LogisticRegression()

        classifier.fit(X_train, y_train)
        y_pred = classifier.score(X_train, y_train)
        classifier.fit(X_test, y_test)
        y_pred1 = classifier.score(X_test, y_test)
        score_total+=y_pred1
        tab[valeur].set(round(y_pred1,4))

        #print('Accuracy {} data: {:.4f}'.format(file,y_pred))
        #print('Accuracy {} data: {}'.format(file,y_pred1))
        valeur+=1
    return score_total/4
word_list={}
def verify_other_factor(doc,word):
    score=0
    f=open('certificat_name.txt')

    for ele in other:
        cert = ""
        value = 0
        for para in doc.paragraphs:

            if para.text != "" and para.text[0].isdigit():
                if ele in para.text:
                    value = 1
                else:
                    value = 2
            elif value == 1:
                cert += para.text
        word_list[ele] = cert
    for ele in word_list.keys():
        print(ele+': '+word_list[ele])
    if '2048 bit' in word_list['KEY SIZE']:
        score=0.8
    if 'The authentication of a candidate Isabel Customer’s identity' in word_list['AUTHENTICATION OF ORGANIZATION IDENTITY']:
        score+=0.7
    for ele in f.readlines():
        if ele in word["CERTIFICATION AUTHORITIES"]:
            score+=1
    Texte8.set(round(score/3,2))
    return score/3
def test(document):
    score_total=0
    document=Document(document)
    with open('file.txt', "w", encoding="utf-8") as textFile:

        var=""
        for ele in principale:
            cert = ""
            value=0
            for para in document.paragraphs:

                if para.text!="" and para.text[0].isdigit()  :
                    if ele in para.text:
                        value=1
                    else:
                        value=2
                elif value==1:
                    cert += para.text
            word[ele]=cert



    for ele in word.keys():
        print(ele+': '+word[ele])
    valeur=0
    for key in word.keys():
        list = []
        fichier=key.split(' ')
        file=fichier[0]+'_'+fichier[1]+'.txt'
        print(file)
        dat = pd.read_csv(file, names=['sentence'], sep='_')
        list.append(dat)
        df = pd.concat(list)
        print("--------------")
        print(df['sentence'].values)

        score=0
        string = []
        for line in df['sentence'].values:
            tokens = wordpunct_tokenize(line)
            print(tokens)
            for ele in tokens:
                string +=[ele]
        string=set(string)


        field=wordpunct_tokenize(word[key])
        for ele in string:
                if ele in field:
                    score+=1

        score_total+=score/len(string)
        tab1[valeur].set(round(score/len(string),4))
        print(score,len(string),scoree)
        valeur+=1
    print('####################################')
    val1=model_train(word)
    val2=verify_other_factor(document,word)
    Texte9.set(round((val1+val2+score_total/4)/3,4))







afficher=Label(fenetre, text = "SVP votre PC : ",width=120)
afficher.pack()
sais=Text(fenetre,width=60,height=10)
sais.pack()


def UploadAction(event=None):
    filename = filedialog.askopenfilename(filetypes = (("pdf or word files","*.docx"),("all files","*.*")))
    sais.insert(INSERT,filename)
def valider():
    valeur = sais.get("1.0", "end-1c")
    test(valeur)


button = Button(fenetre, text='Open', command=UploadAction)
button.pack()
button1 = Button(fenetre, text='valider', command=valider)
button1.pack()
afficher1=Label(fenetre,text = "CERTIFICATION AUTH %",width=25)
afficher1.place(x=0,y=240)
afficher2=Label(fenetre, textvariable = Texte0,width=5,bg='yellow')
afficher2.place(x=160,y=240)
afficher3=Label(fenetre, textvariable = Texte1,width=5,bg='green')
afficher3.place(x=200,y=240)
afficher1=Label(fenetre,text = "REGISTRATION AUTH %",width=25)
afficher1.place(x=240,y=240)
afficher2=Label(fenetre, textvariable = Texte2,width=5,bg='yellow')
afficher2.place(x=400,y=240)
afficher3=Label(fenetre, textvariable = Texte3,width=5,bg='green')
afficher3.place(x=440,y=240)

afficher1=Label(fenetre,text = "VALIDATION AUTH %",width=25)
afficher1.place(x=0,y=270)
afficher2=Label(fenetre, textvariable = Texte4,width=5,bg='yellow')
afficher2.place(x=160,y=270)
afficher3=Label(fenetre, textvariable = Texte5,width=5,bg='green')
afficher3.place(x=200,y=270)
afficher1=Label(fenetre,text = "CERTIFICATE ACCEPTANCE %",width=25)
afficher1.place(x=240,y=270)
afficher2=Label(fenetre, textvariable = Texte6,width=5,bg='yellow')
afficher2.place(x=400,y=270)
afficher3=Label(fenetre, textvariable = Texte7,width=5,bg='green')
afficher3.place(x=440,y=270)
afficher1=Label(fenetre,text = "OTHER FACTOR %",width=25)
afficher1.place(x=100,y=300)
afficher2=Label(fenetre, textvariable = Texte8,width=5,bg='yellow')
afficher2.place(x=270,y=300)
afficher1=Label(fenetre,text = "VALIDATION FINALE %",width=25)
afficher1.place(x=200,y=330)
afficher2=Label(fenetre, textvariable = Texte9,width=10,bg='green')
afficher2.place(x=370,y=330)

fenetre.mainloop()



